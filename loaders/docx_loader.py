#!/usr/bin/env python3
"""
DOCX Loader with Multi-Modal Extraction and OpenTelemetry Instrumentation

Enhanced with comprehensive OpenTelemetry observability following EDOT best practices
for distributed tracing, metrics, and structured logging correlation.

Demonstrates complete DOCX processing observability WITHOUT LOGSTASH DEPENDENCY - showing how OpenTelemetry 
can handle multiple log types and correlate them with trace IDs and span IDs for ELK visualization.
"""

import os
import time
import socket
import uuid
import logging
from typing import List, Dict, Any

# OpenTelemetry imports - Official EDOT libraries
from opentelemetry import trace, metrics, baggage
from opentelemetry.instrumentation.logging import LoggingInstrumentor
from opentelemetry.sdk.trace import TracerProvider
from opentelemetry.sdk.trace.export import BatchSpanProcessor, ConsoleSpanExporter
from opentelemetry.sdk.metrics import MeterProvider
from opentelemetry.sdk.metrics.export import PeriodicExportingMetricReader, ConsoleMetricExporter
from opentelemetry.sdk.resources import Resource, SERVICE_NAME, SERVICE_VERSION, SERVICE_INSTANCE_ID
from opentelemetry.exporter.otlp.proto.grpc.trace_exporter import OTLPSpanExporter
from opentelemetry.exporter.otlp.proto.grpc.metric_exporter import OTLPMetricExporter
from opentelemetry.propagate import set_global_textmap, extract
from opentelemetry.propagators.b3 import B3MultiFormat

from langchain_core.documents import Document
from docx import Document as DocxDocument
import mammoth
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from loaders.image_extractor import extract_images_from_docx
from tools.vision_tool import analyze_images_with_vision_model, VISION_PROMPT

# OCR imports
try:
    import easyocr
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: easyocr not installed. OCR text extraction will be skipped.")

load_dotenv()

# Initialize OpenTelemetry with EDOT configuration
def init_telemetry():
    """Initialize OpenTelemetry with EDOT configuration and context extraction"""
    
    # Check if we're already using auto-instrumentation
    current_tracer_provider = trace.get_tracer_provider()
    if hasattr(current_tracer_provider, '__class__') and current_tracer_provider.__class__.__name__ != 'ProxyTracerProvider':
        print("✅ Using existing OpenTelemetry auto-instrumentation")
        return trace.get_tracer(__name__), metrics.get_meter(__name__)
    
    # Extract trace context from environment variables set by orchestrator
    carrier = {}
    for key, value in os.environ.items():
        if key.startswith('OTEL_PROPAGATED_'):
            header_key = key.replace('OTEL_PROPAGATED_', '').replace('_', '-').lower()
            carrier[header_key] = value
    
    # Extract and set context from orchestrator
    extracted_context = extract(carrier) if carrier else None
    if extracted_context:
        print("✅ Extracted trace context from orchestrator")
    
    # Generate unique service instance ID
    service_instance_id = f"{socket.gethostname()}-{uuid.uuid4().hex[:8]}"
    
    # Resource configuration - Following EDOT best practices
    resource = Resource.create({
        SERVICE_NAME: "document-rag-docx-loader",
        SERVICE_VERSION: "1.0.0",
        SERVICE_INSTANCE_ID: service_instance_id,
        "service.namespace": "document-rag-system",
        "deployment.environment": os.getenv("DEPLOYMENT_ENV", "production"),
        "host.name": socket.gethostname(),
        "process.pid": str(os.getpid()),
        "telemetry.sdk.language": "python",
        "telemetry.sdk.name": "opentelemetry",
        "telemetry.sdk.version": "1.25.0"
    })
    
    # OTLP endpoint configuration
    otlp_endpoint = os.getenv("OTEL_EXPORTER_OTLP_ENDPOINT", "http://localhost:4317")
    
    try:
        # Trace provider setup
        trace_provider = TracerProvider(resource=resource)
        trace.set_tracer_provider(trace_provider)
        
        # OTLP span exporter
        otlp_span_exporter = OTLPSpanExporter(
            endpoint=otlp_endpoint,
            insecure=True,
            headers={}
        )
        
        # Add span processors
        trace_provider.add_span_processor(BatchSpanProcessor(otlp_span_exporter))
        
        # Console exporter for debugging
        if os.getenv("OTEL_DEBUG", "false").lower() == "true":
            trace_provider.add_span_processor(BatchSpanProcessor(ConsoleSpanExporter()))
        
        # Metrics provider setup
        otlp_metric_exporter = OTLPMetricExporter(
            endpoint=otlp_endpoint,
            insecure=True,
            headers={}
        )
        
        metric_reader = PeriodicExportingMetricReader(
            exporter=otlp_metric_exporter,
            export_interval_millis=15000,  # 15 seconds
        )
        
        metric_provider = MeterProvider(
            resource=resource,
            metric_readers=[metric_reader]
        )
        metrics.set_meter_provider(metric_provider)
        
        print("✅ OpenTelemetry initialized for DOCX loader")
        
    except Exception as e:
        print(f"⚠️  OpenTelemetry provider setup error: {e}")
    
    # Set up B3 propagator for distributed tracing
    set_global_textmap(B3MultiFormat())
    
    # Manual instrumentation only if not using auto-instrumentation
    if not os.getenv("OTEL_PYTHON_DISABLED_INSTRUMENTATIONS"):
        try:
            LoggingInstrumentor().instrument(set_logging_format=True)
        except Exception as e:
            print(f"⚠️  Instrumentation warning: {e}")
    
    return trace.get_tracer(__name__), metrics.get_meter(__name__)

# Initialize telemetry
tracer, meter = init_telemetry()

# Custom metrics for DOCX processing observability
docx_processing_counter = meter.create_counter(
    "docx_documents_processed_total",
    description="Total number of DOCX documents processed",
    unit="1"
)

docx_processing_duration = meter.create_histogram(
    "docx_processing_duration_seconds",
    description="Time taken to process DOCX documents",
    unit="s"
)

docx_content_extraction_counter = meter.create_counter(
    "docx_content_extraction_total",
    description="Total number of content extractions by type",
    unit="1"
)

docx_images_counter = meter.create_counter(
    "docx_images_extracted_total",
    description="Total number of images extracted from DOCX",
    unit="1"
)

mammoth_conversion_counter = meter.create_counter(
    "mammoth_conversions_total",
    description="Total number of Mammoth HTML conversions",
    unit="1"
)

ocr_operations_counter = meter.create_counter(
    "ocr_operations_total",
    description="Total number of OCR operations performed",
    unit="1"
)

vision_analysis_counter = meter.create_counter(
    "vision_analysis_operations_total",
    description="Total number of vision AI analysis operations",
    unit="1"
)

# Configure structured logging with OpenTelemetry correlation
class OtelDocxLoaderFormatter(logging.Formatter):
    """Custom formatter to include OpenTelemetry trace and span IDs"""
    
    def format(self, record):
        # Get current span context for correlation
        span = trace.get_current_span()
        if span.get_span_context().is_valid:
            trace_id = format(span.get_span_context().trace_id, '032x')
            span_id = format(span.get_span_context().span_id, '016x')
            record.trace_id = trace_id
            record.span_id = span_id
        else:
            record.trace_id = '00000000000000000000000000000000'
            record.span_id = '0000000000000000'
        
        return super().format(record)

# Set up structured logging
formatter = OtelDocxLoaderFormatter(
    '%(asctime)s - %(name)s - %(levelname)s - [trace_id=%(trace_id)s span_id=%(span_id)s] - %(message)s'
)

console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)

file_handler = logging.FileHandler('docx_loader.log')
file_handler.setFormatter(formatter)

logging.basicConfig(
    level=logging.INFO,
    handlers=[console_handler, file_handler]
)
logger = logging.getLogger(__name__)

DOCX_IMAGE_DIR = "docx_images"
os.makedirs(DOCX_IMAGE_DIR, exist_ok=True)

def extract_text_with_ocr(image_path: str) -> str:
    """Extract text from a DOCX image using OCR with comprehensive instrumentation."""
    with tracer.start_as_current_span("extract_text_with_ocr") as span:
        span.set_attribute("image.path", image_path[:100])  # Truncate for security
        span.set_attribute("ocr.available", OCR_AVAILABLE)
        
        ocr_operations_counter.add(1, {"status": "attempt"})
        
        if not OCR_AVAILABLE:
            span.set_attribute("ocr.skipped", True)
            span.add_event("OCR not available, skipping")
            return ""

        try:
            with tracer.start_as_current_span("easyocr_processing") as ocr_span:
                reader = easyocr.Reader(['en'])
                results = reader.readtext(image_path)
                
                ocr_span.set_attribute("ocr.results_count", len(results))

                # Combine all detected text
                ocr_text = []
                high_confidence_count = 0
                
                for (bbox, text, confidence) in results:
                    if confidence > 0.5:  # Only include high-confidence text
                        ocr_text.append(text)
                        high_confidence_count += 1

                final_text = "\n".join(ocr_text)
                
                ocr_span.set_attribute("ocr.high_confidence_results", high_confidence_count)
                ocr_span.set_attribute("ocr.extracted_text_length", len(final_text))
                ocr_span.add_event("OCR text extraction completed")

            ocr_operations_counter.add(1, {"status": "success"})
            
            span.set_attribute("ocr.extraction_successful", True)
            span.set_attribute("ocr.text_length", len(final_text))
            span.add_event("OCR extraction completed successfully")
            
            logger.info(f"✅ OCR extracted {len(final_text)} characters from {os.path.basename(image_path)}")
            
            return final_text

        except Exception as e:
            span.record_exception(e)
            span.set_attribute("ocr.extraction_successful", False)
            
            ocr_operations_counter.add(1, {"status": "error"})
            
            logger.error(f"❌ OCR extraction failed for {image_path}: {e}")
            return ""

def extract_tables(doc: DocxDocument) -> List[Dict[str, Any]]:
    """Extract tables from the document with their structure preserved and instrumentation."""
    with tracer.start_as_current_span("extract_docx_tables") as span:
        tables = []
        
        span.set_attribute("tables.found", len(doc.tables))
        
        for i, table in enumerate(doc.tables):
            with tracer.start_as_current_span(f"process_table_{i+1}") as table_span:
                # Get header row
                header = []
                for cell in table.rows[0].cells:
                    header.append(cell.text.strip())

                # Get data rows
                data = []
                for row in table.rows[1:]:
                    row_data = {}
                    for j, cell in enumerate(row.cells):
                        if j < len(header):  # Ensure we have a header for this column
                            row_data[header[j]] = cell.text.strip()
                    data.append(row_data)

                table_info = {
                    "header": header,
                    "data": data
                }
                tables.append(table_info)
                
                table_span.set_attribute("table.columns", len(header))
                table_span.set_attribute("table.rows", len(data))
                table_span.add_event("Table processed successfully")

        span.set_attribute("tables.extracted", len(tables))
        span.add_event("Table extraction completed", {"tables_count": len(tables)})
        
        docx_content_extraction_counter.add(len(tables), {"content_type": "table"})
        
        logger.info(f"📊 Extracted {len(tables)} tables from DOCX")
        return tables

def extract_structured_content(doc: DocxDocument) -> List[Dict[str, Any]]:
    """Extract content with structure (paragraphs, headings, lists) preserved and instrumentation."""
    with tracer.start_as_current_span("extract_structured_content") as span:
        content = []
        content_type_counts = {}

        for paragraph in doc.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue

            # Determine content type based on style
            style = paragraph.style.name.lower()
            content_type = "paragraph"
            if "heading" in style:
                content_type = f"heading{style[-1]}" if style[-1].isdigit() else "heading"
            elif "list" in style:
                content_type = "list_item"

            content_type_counts[content_type] = content_type_counts.get(content_type, 0) + 1

            content.append({
                "type": content_type,
                "text": paragraph.text.strip(),
                "style": style
            })

        span.set_attribute("content.paragraphs", len(content))
        span.set_attribute("content.type_variety", len(content_type_counts))
        
        # Record metrics for each content type
        for content_type, count in content_type_counts.items():
            docx_content_extraction_counter.add(count, {"content_type": content_type})

        span.add_event("Structured content extraction completed", {
            "paragraphs": len(content),
            "content_types": list(content_type_counts.keys())
        })
        
        logger.info(f"📝 Extracted {len(content)} paragraphs with {len(content_type_counts)} content types")
        return content

def extract_semantic_html(file_path: str) -> str:
    """Convert DOCX to semantic HTML using Mammoth with instrumentation."""
    with tracer.start_as_current_span("extract_semantic_html") as span:
        span.set_attribute("file.path", file_path[:100])  # Truncate for security
        
        mammoth_conversion_counter.add(1, {"status": "attempt"})
        
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                
                span.set_attribute("html.length", len(html_content))
                span.set_attribute("conversion.successful", True)
                span.add_event("Mammoth conversion completed successfully")
                
                mammoth_conversion_counter.add(1, {"status": "success"})
                
                logger.info(f"🔄 Converted DOCX to HTML ({len(html_content)} characters)")
                return html_content
                
        except Exception as e:
            span.record_exception(e)
            span.set_attribute("conversion.successful", False)
            
            mammoth_conversion_counter.add(1, {"status": "error"})
            
            logger.error(f"❌ Mammoth conversion failed: {e}")
            raise

def parse_html_content(html: str) -> List[Dict[str, Any]]:
    """Parse HTML content into structured blocks with instrumentation."""
    with tracer.start_as_current_span("parse_html_content") as span:
        span.set_attribute("html.length", len(html))
        
        soup = BeautifulSoup(html, 'html.parser')
        content = []
        element_counts = {}

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'table']):
            element_type = element.name
            element_counts[element_type] = element_counts.get(element_type, 0) + 1
            
            if element.name.startswith('h'):
                content.append({
                    "type": f"heading{element.name[1]}",
                    "text": element.get_text().strip()
                })
            elif element.name == 'p':
                content.append({
                    "type": "paragraph",
                    "text": element.get_text().strip()
                })
            elif element.name in ['ul', 'ol']:
                list_type = "unordered" if element.name == 'ul' else "ordered"
                items = [item.get_text().strip() for item in element.find_all('li')]
                content.append({
                    "type": "list",
                    "list_type": list_type,
                    "items": items
                })
            elif element.name == 'table':
                rows = []
                for row in element.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    rows.append(cells)
                content.append({
                    "type": "table",
                    "rows": rows
                })

        span.set_attribute("content.elements_parsed", len(content))
        span.set_attribute("html.element_types", list(element_counts.keys()))
        
        span.add_event("HTML parsing completed", {
            "elements_parsed": len(content),
            "element_types": element_counts
        })
        
        logger.info(f"📄 Parsed HTML content into {len(content)} structured blocks")
        return content

def load_docx_as_langchain_docs(file_path: str) -> List[Document]:
    """
    Load a DOCX file and return a list of LangChain Document objects with comprehensive text extraction and instrumentation.

    Uses multiple extraction methods:
    1. Traditional text extraction via python-docx
    2. Semantic HTML extraction via Mammoth
    3. OCR on extracted images
    4. Vision AI analysis for context

    Args:
        file_path (str): Path to the DOCX file to process.

    Returns:
        List[Document]: List of LangChain Document objects containing all extractable content.
    """
    with tracer.start_as_current_span("load_docx_as_langchain_docs") as span:
        start_time = time.time()
        
        span.set_attribute("docx.file_path", file_path[:100])  # Truncate for security
        span.set_attribute("docx.filename", os.path.basename(file_path))
        
        # Add baggage for cross-service correlation
        baggage.set_baggage("docx.filename", os.path.basename(file_path))
        
        docx_processing_counter.add(1, {"status": "attempt"})
        
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                span.set_attribute("docx.file_exists", False)
                docx_processing_counter.add(1, {"status": "error", "error_type": "file_not_found"})
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            span.set_attribute("docx.file_exists", True)
            span.set_attribute("docx.file_size", os.path.getsize(file_path))
            
            docs = []

            # 1. Extract structured content using python-docx
            with tracer.start_as_current_span("extract_structured_content_docx") as extract_span:
                extract_span.set_attribute("extraction.method", "python_docx")
                
                logger.info(f"🔄 Starting DOCX structured content extraction for {os.path.basename(file_path)}")
                
                doc = DocxDocument(file_path)
                structured_content = extract_structured_content(doc)
                tables = extract_tables(doc)
                
                extract_span.set_attribute("structured_content.count", len(structured_content))
                extract_span.set_attribute("tables.count", len(tables))
                extract_span.add_event("Python-docx extraction completed")

            # 2. Extract semantic HTML using Mammoth
            with tracer.start_as_current_span("extract_semantic_html_content") as html_span:
                html_span.set_attribute("extraction.method", "mammoth")
                
                logger.info(f"🔄 Converting DOCX to semantic HTML...")
                
                html_content = extract_semantic_html(file_path)
                semantic_content = parse_html_content(html_content)
                
                html_span.set_attribute("html.content_length", len(html_content))
                html_span.set_attribute("semantic_content.count", len(semantic_content))
                html_span.add_event("Mammoth HTML extraction completed")

            # 3. Create documents from structured content
            with tracer.start_as_current_span("create_structured_documents") as struct_span:
                structured_docs_created = 0
                
                for content in structured_content:
                    docs.append(Document(
                        page_content=content["text"],
                        metadata={
                            "source": file_path,
                            "file_type": "docx",
                            "content_type": content["type"],
                            "style": content.get("style", "default"),
                            "extraction_methods": "python_docx,mammoth"
                        }
                    ))
                    structured_docs_created += 1

                struct_span.set_attribute("structured_docs.created", structured_docs_created)
                struct_span.add_event("Structured documents created")

            # 4. Create documents from tables
            with tracer.start_as_current_span("create_table_documents") as table_span:
                table_docs_created = 0
                
                for table in tables:
                    # Convert table to string representation
                    table_str = "\n".join([
                        "\t".join(table["header"]),  # Header row
                        *["\t".join(str(val) for val in row.values()) for row in table["data"]]  # Data rows
                    ])
                    docs.append(Document(
                        page_content=table_str,
                        metadata={
                            "source": file_path,
                            "file_type": "docx",
                            "content_type": "table",
                            "columns": table["header"],
                            "extraction_methods": "python_docx"
                        }
                    ))
                    table_docs_created += 1

                table_span.set_attribute("table_docs.created", table_docs_created)
                table_span.add_event("Table documents created")

            # 5. Extract and analyze images with comprehensive approach
            images = []
            with tracer.start_as_current_span("extract_docx_images") as img_extract_span:
                logger.info(f"🖼️ Extracting images from DOCX...")
                
                images = extract_images_from_docx(file_path, DOCX_IMAGE_DIR)
                
                img_extract_span.set_attribute("images.extracted_count", len(images))
                img_extract_span.add_event("Image extraction completed")
                
                docx_images_counter.add(len(images), {"source": "docx_extraction"})
                
                logger.info(f"🖼️ Extracted {len(images)} images from DOCX")

            # Enhanced Vision AI analysis with comprehensive prompt
            image_docs = []
            if images:
                with tracer.start_as_current_span("vision_ai_analysis") as vision_span:
                    vision_prompt = """
                    Analyze this DOCX image and extract:
                    1. ALL visible text (exact transcription)
                    2. Key concepts and information
                    3. Context and meaning
                    4. Any structured data (lists, tables, diagrams, charts)
                    5. Technical diagrams or flowcharts
                    6. Any text that might be in images or graphics

                    If this appears to be a document page or embedded content, transcribe all text exactly as it appears.
                    Focus on extracting actionable information that would be useful for answering questions.
                    Include any text that might be embedded in images, charts, or diagrams.
                    """
                    
                    vision_span.set_attribute("analysis.method", "vision_ai")
                    vision_span.set_attribute("images.to_analyze", len(images))
                    vision_span.set_attribute("vision.prompt_length", len(vision_prompt))
                    
                    vision_analysis_counter.add(1, {"status": "attempt", "images_count": len(images)})
                    
                    logger.info(f"🤖 Analyzing {len(images)} images with Vision AI...")
                    
                    image_docs = analyze_images_with_vision_model(images, prompt=vision_prompt)
                    
                    vision_span.set_attribute("vision_docs.generated", len(image_docs))
                    vision_span.add_event("Vision AI analysis completed")
                    
                    vision_analysis_counter.add(1, {"status": "success", "docs_generated": len(image_docs)})
                    
                    logger.info(f"🤖 Vision AI generated {len(image_docs)} document analyses")

            # 6. Add OCR text extraction to image documents
            with tracer.start_as_current_span("ocr_text_enhancement") as ocr_span:
                ocr_enhanced_count = 0
                total_ocr_text_length = 0
                
                for img_doc in image_docs:
                    img_path = img_doc.metadata.get("image_path")
                    if img_path and os.path.exists(img_path):
                        with tracer.start_as_current_span(f"ocr_process_image") as img_ocr_span:
                            img_ocr_span.set_attribute("image.path", os.path.basename(img_path))
                            
                            ocr_text = extract_text_with_ocr(img_path)
                            if ocr_text:
                                # Combine vision AI analysis with OCR text
                                combined_content = f"{img_doc.page_content}\n\n[OCR Extracted Text]\n{ocr_text}"
                                img_doc.page_content = combined_content
                                img_doc.metadata["extraction_methods"] = "vision_ai,ocr"
                                ocr_enhanced_count += 1
                                total_ocr_text_length += len(ocr_text)
                                
                                img_ocr_span.set_attribute("ocr.text_added", True)
                                img_ocr_span.set_attribute("ocr.text_length", len(ocr_text))
                            else:
                                img_doc.metadata["extraction_methods"] = "vision_ai"
                                img_ocr_span.set_attribute("ocr.text_added", False)

                ocr_span.set_attribute("ocr.enhanced_documents", ocr_enhanced_count)
                ocr_span.set_attribute("ocr.total_text_length", total_ocr_text_length)
                ocr_span.add_event("OCR enhancement completed")

            docs.extend(image_docs)

            # Calculate processing metrics
            processing_duration = time.time() - start_time
            total_docs = len(docs)
            
            docx_processing_duration.record(processing_duration, {
                "docx_size": "large" if os.path.getsize(file_path) > 5_000_000 else "small",
                "status": "success"
            })
            
            docx_processing_counter.add(1, {"status": "success"})

            # Update span with final metrics
            span.set_attribute("docx.processing_duration_seconds", processing_duration)
            span.set_attribute("docx.total_documents_generated", total_docs)
            span.set_attribute("docx.structured_documents", len(structured_content))
            span.set_attribute("docx.table_documents", len(tables))
            span.set_attribute("docx.image_documents", len(image_docs))
            span.set_attribute("docx.images_extracted", len(images))
            span.set_attribute("docx.processing_successful", True)
            
            span.add_event("DOCX processing completed successfully", {
                "total_documents": total_docs,
                "structured_documents": len(structured_content),
                "table_documents": len(tables),
                "image_documents": len(image_docs),
                "images_extracted": len(images),
                "processing_duration": processing_duration
            })

            logger.info(f"✅ DOCX processing completed: {total_docs} total documents in {processing_duration:.2f}s")

            return docs

        except Exception as e:
            processing_duration = time.time() - start_time
            
            span.record_exception(e)
            span.set_attribute("docx.processing_successful", False)
            span.set_attribute("docx.processing_duration_seconds", processing_duration)
            span.set_attribute("error.type", type(e).__name__)
            
            docx_processing_counter.add(1, {"status": "error", "error_type": type(e).__name__})
            
            docx_processing_duration.record(processing_duration, {
                "docx_size": "unknown",
                "status": "error"
            })
            
            logger.error(f"❌ DOCX processing failed for {file_path}: {e}")
            raise

if __name__ == "__main__":
    """Example usage with OpenTelemetry instrumentation"""
    import sys
    
    with tracer.start_as_current_span("docx_loader_main") as span:
        span.set_attribute("execution.mode", "standalone")
        
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            span.set_attribute("file.path", file_path)
            
            logger.info(f"🧪 Testing DOCX loader: {file_path}")
            
            try:
                docs = load_docx_as_langchain_docs(file_path)
                
                span.set_attribute("documents.loaded", len(docs))
                span.add_event("DOCX loading test completed successfully")
                
                logger.info(f"📊 Test Results: Processed {len(docs)} documents from {file_path}")

                # Group documents by type for better overview
                content_types = {}
                for doc in docs:
                    content_type = doc.metadata.get("content_type", "unknown")
                    if content_type not in content_types:
                        content_types[content_type] = 0
                    content_types[content_type] += 1

                logger.info("\n📋 Content breakdown:")
                for content_type, count in content_types.items():
                    logger.info(f"  - {content_type}: {count} items")

                logger.info("\n📄 First few documents preview:")
                for i, doc in enumerate(docs[:3], 1):
                    logger.info(f"\nDocument {i}:")
                    logger.info(f"  Type: {doc.metadata.get('content_type', 'unknown')}")
                    logger.info(f"  Content preview: {doc.page_content[:200]}...")
                    
            except Exception as e:
                span.record_exception(e)
                span.set_attribute("test.successful", False)
                logger.error(f"❌ Test failed: {e}")
                raise
        else:
            span.set_attribute("usage", "help")
            logger.info("Usage: python docx_loader.py <path_to_docx_file>")

